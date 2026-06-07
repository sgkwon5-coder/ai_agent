"""Build the Coca-Cola Company corporate research report (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_PATH = r"C:\agent_k\Coca-Cola_Company_Report.docx"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.3

# ---------------------------------------------------------------------------
# Cover / Title
# ---------------------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("The Coca-Cola Company")
title_run.font.size = Pt(28)
title_run.font.bold = True

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
title2_run = title2.add_run("Corporate Research Report")
title2_run.font.size = Pt(18)
title2_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run(
    "Company Overview, Products, Financial Highlights, and Global Strategy"
)
subtitle_run.font.size = Pt(13)
subtitle_run.font.italic = True

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run("June 7, 2026")
date_run.font.size = Pt(12)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Helper to build simple two-column tables
# ---------------------------------------------------------------------------

def add_table(headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value

    if col_widths_cm:
        for i, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


# ---------------------------------------------------------------------------
# 1. Company Overview
# ---------------------------------------------------------------------------
doc.add_heading("1. Company Overview", level=1)

overview_items = [
    ("Founded", "Formula developed in May 1886 by pharmacist John Stith "
                "Pemberton; The Coca-Cola Company was incorporated in 1892."),
    ("Headquarters", "Atlanta, Georgia, USA"),
    ("Industry", "Food, Beverages & Tobacco / Nonalcoholic Beverage Manufacturing"),
    ("Stock", "Listed on the New York Stock Exchange under the ticker \"KO\"; "
              "a component of the Dow Jones Industrial Average (DJIA), S&P 500, "
              "and S&P 100 indices."),
    ("Leadership", "James Quincey currently serves as CEO and is scheduled to "
                   "transition to the role of Executive Chairman effective "
                   "March 31, 2026. Henrique Braun, currently Executive Vice "
                   "President and Chief Operating Officer, will become CEO at "
                   "that time."),
    ("Scale", "More than 500 brands sold in over 200 countries and "
              "territories; one of the world's largest beverage manufacturers "
              "and distributors."),
    ("Operating Segments", "Europe, Middle East & Africa; Latin America; "
                           "North America; Asia Pacific; Bottling Investments; "
                           "Global Ventures."),
]

for label, text in overview_items:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(text)

# ---------------------------------------------------------------------------
# 2. Core Products and Brand Portfolio
# ---------------------------------------------------------------------------
doc.add_heading("2. Core Products and Brand Portfolio", level=1)

doc.add_paragraph(
    "The Coca-Cola Company's portfolio spans several beverage categories. "
    "Representative brands within each category are summarized below."
)

add_table(
    headers=["Category", "Representative Brands"],
    rows=[
        ("Sparkling Soft Drinks", "Coca-Cola, Sprite, Fanta"),
        ("Water, Sports, Coffee & Tea",
         "Dasani, smartwater, vitaminwater, Topo Chico, BODYARMOR, Powerade, "
         "Costa, Georgia, Fuze Tea, Gold Peak, Ayataka"),
        ("Juice, Dairy & Plant-Based Beverages",
         "Minute Maid, Simply, innocent, Del Valle, fairlife, AdeS"),
    ],
    col_widths_cm=[5, 11],
)

doc.add_paragraph()
note_p = doc.add_paragraph()
note_p.add_run("Note: ").bold = True
note_p.add_run(
    "BODYARMOR expanded into Europe (Spain) in 2025 as part of an "
    "international expansion strategy that builds on its established position "
    "in the U.S. sports drink category. The fairlife portfolio continues to "
    "expand within the U.S. dairy segment."
)

# ---------------------------------------------------------------------------
# 3. Financial Highlights — Full Year 2025
# ---------------------------------------------------------------------------
doc.add_heading("3. Financial Highlights — Full Year 2025", level=1)

add_table(
    headers=["Metric", "Value / Result"],
    rows=[
        ("Net Revenues",
         "$47.9 billion (up 2% year-over-year; organic revenue growth of 5%)"),
        ("Full-Year EPS",
         "$3.04 (up 23%, driven by global value share gains in nonalcoholic "
         "ready-to-drink beverages and emerging-market performance in Brazil "
         "and Central Asia)"),
        ("Operating Income",
         "Full-year operating income grew 38%; comparable currency-neutral "
         "operating income grew 13% on cost management and pricing actions. "
         "Q4 operating income declined 32% due to a $960 million non-cash "
         "impairment charge related to the BODYARMOR trademark."),
        ("Cash Flow from Operations", "$7.4 billion (full year)"),
        ("Free Cash Flow (Non-GAAP)", "$5.3 billion (full year)"),
        ("2026 Outlook",
         "Company guidance projects organic revenue growth of 4%–5% and "
         "comparable EPS growth of 7%–8%"),
    ],
    col_widths_cm=[5, 11],
)

doc.add_paragraph()
fin_note = doc.add_paragraph()
fin_note.add_run(
    "These figures are drawn from the company's fourth-quarter and full-year "
    "2025 earnings release (reported in February 2026) and are subject to "
    "revision in subsequent SEC filings."
).italic = True

# ---------------------------------------------------------------------------
# 4. Global Strategy and Outlook
# ---------------------------------------------------------------------------
doc.add_heading("4. Global Strategy and Outlook", level=1)

strategy_items = [
    ("Digital Transformation",
     "The company created a new Chief Digital Officer role to unify digital, "
     "data, and operational excellence across the organization, aiming to "
     "improve speed and efficiency from consumer engagement to internal "
     "operations."),
    ("Innovation Infrastructure",
     "Plans call for establishing innovation hubs and commercial centers of "
     "excellence across all operating segments and key markets in 2026."),
    ("Sustainability (\"World Without Waste\")",
     "A long-running packaging and recycling initiative. The company has "
     "revised its voluntary environmental packaging-collection goal "
     "timelines, extending target dates from 2025/2030 to 2035."),
    ("Leadership Transition",
     "Incoming CEO Henrique Braun, who succeeds James Quincey (moving to "
     "Executive Chairman), has signaled a focus on accelerating innovation "
     "and sustaining the company's growth momentum."),
    ("Growth Trajectory",
     "Management points to 5% organic revenue growth in 2025 as evidence of "
     "a \"durable strategy\" with a \"long runway\" for continued expansion."),
]

for label, text in strategy_items:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(text)

# ---------------------------------------------------------------------------
# 5. Sources / References
# ---------------------------------------------------------------------------
doc.add_heading("5. Sources / References", level=1)

sources = [
    "https://investors.coca-colacompany.com/news-events/press-releases/detail/1151/coca-cola-reports-fourth-quarter-and-full-year-2025-results",
    "https://www.foodnavigator.com/Article/2026/01/16/cocacolas-pivotal-year-portfolio-shifts-and-a-new-ceo/",
    "https://www.beveragedaily.com/Article/2026/02/16/cocacola-maps-bold-growth-strategy-for-2026/",
    "https://www.packworld.com/sustainable-packaging/recycling/article/22928147/coca-cola-extends-sustainable-packaging-deadlines-to-2035",
    "https://en.wikipedia.org/wiki/The_Coca-Cola_Company",
]

for src in sources:
    doc.add_paragraph(src, style="List Bullet")

doc.save(OUTPUT_PATH)
print(f"Saved report to: {OUTPUT_PATH}")
